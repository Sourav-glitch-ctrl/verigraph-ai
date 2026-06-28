import logging
from pathlib import Path
from typing import List

from ingestion.models import Document

# pyrefly: ignore [missing-import]
from docx import Document as DocxDocument
# pyrefly: ignore [missing-import]
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

class DOCXLoader:
    """
    Responsible only for extracting text
    from DOCX documents.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        doc_path = Path(self.file_path)

        if not doc_path.exists():
            logger.error(f"DOCX file not found: {self.file_path}")
            raise FileNotFoundError(f"DOCX file not found: {self.file_path}")

        if not doc_path.is_file():
            logger.error(f"Path is not a file: {self.file_path}")
            raise ValueError(f"Path is not a file: {self.file_path}")

        if doc_path.suffix.lower() != '.docx':
            logger.warning(f"File {self.file_path} does not have a .docx extension.")

        docs = []
        try:
            document = DocxDocument(self.file_path)
            paragraphs = []
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            if paragraphs:
                full_text = "\n".join(paragraphs)
                docs.append(Document(
                    page_content=full_text,
                    metadata={"source": str(doc_path)}
                ))
                
        except PackageNotFoundError as e:
            logger.error(f"Invalid or corrupted DOCX file {self.file_path}: {e}")
            raise ValueError(f"Invalid or corrupted DOCX file: {self.file_path}") from e
        except Exception as e:
            logger.error(f"Unexpected error while loading DOCX {self.file_path}: {e}")
            raise RuntimeError(f"Unexpected error while loading DOCX: {self.file_path}") from e

        return docs